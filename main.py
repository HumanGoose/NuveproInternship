import tkinter as tk
import docx
from cryptography.fernet import Fernet

# Generate a secret key for encryption. You should securely store this key.
# You can generate one using `Fernet.generate_key()` and store it securely.
# For simplicity, we'll generate one here, but it's important to keep it safe in a real application.
SECRET_KEY = Fernet.generate_key()
cipher_suite = Fernet(SECRET_KEY)

def login():
    username = username_entry.get()
    password = password_entry.get()

    # Encrypt the password
    encrypted_password = cipher_suite.encrypt(password.encode())

    # Save the data in a Word document
    save_credentials(username, encrypted_password)

    open_new_window()
    window.withdraw()
    username_entry.delete(0, 'end')
    password_entry.delete(0, 'end')

def save_credentials(username, encrypted_password):
    doc = docx.Document()
    try:
        doc = docx.Document('user_credentials.docx')  # Open the existing document
    except FileNotFoundError:
        pass  # If the document doesn't exist, create a new one

    doc.add_paragraph(f'Username: {username}, Password: {encrypted_password.decode()}')
    doc.save('user_credentials.docx')

def decrypt_password(encrypted_password):
    decrypted_password = cipher_suite.decrypt(encrypted_password.encode())
    return decrypted_password.decode()

# Create a simple login window
window = tk.Tk()
window.title("Login Page")

# Set a fixed size for the window
window.geometry("400x200")

# Center the window on the screen
window.eval('tk::PlaceWindow . center')

# Styling for labels and entry widgets
label_style = {
    'font': ('Arial', 12),
    'padx': 10,
    'pady': 10
}

entry_style = {
    'font': ('Arial', 12),
    'width': 20
}

def open_new_window():
    global new_window  # Use the global variable

    new_window = tk.Toplevel(window)  # Create a new window

    new_window.title("Welcome Page")

    # Set a fixed size for the window
    new_window.geometry("400x200")

    # Center the window on the screen
    new_window.geometry("400x200+{}+{}".format(
        window.winfo_x() + (window.winfo_width() - new_window.winfo_reqwidth()) // 2,
        window.winfo_y() + (window.winfo_height() - new_window.winfo_reqheight()) // 2
    ))

    # Styling for labels
    label_style = {
        'font': ('Arial', 16),
        'padx': 10,
        'pady': 10
    }

    # Welcome label
    welcome_label = tk.Label(new_window, text="Welcome! You have successfully logged in.", **label_style)
    welcome_label.grid(row=0, column=0)

    # Logout button
    logout_button = tk.Button(new_window, text="Logout", command=close_welcome_window, **label_style)
    logout_button.grid(row=1, column=0)

def close_welcome_window():
    window.deiconify()  # Show the login window
    new_window.destroy()  # Close the welcome window

# Username label and entry
username_label = tk.Label(window, text="Username:", **label_style)
username_label.grid(row=0, column=0)
username_entry = tk.Entry(window, **entry_style)
username_entry.grid(row=0, column=1)

# Password label and entry
password_label = tk.Label(window, text="Password:", **label_style)
password_label.grid(row=1, column=0)
password_entry = tk.Entry(window, show="*", **entry_style)  # Passwords are hidden with '*'
password_entry.grid(row=1, column=1)

# Login button
login_button = tk.Button(window, text="Login", command=login, **label_style)
login_button.grid(row=2, column=0, columnspan=2)

# Result label
result_label = tk.Label(window, text="", **label_style)
result_label.grid(row=3, column=0, columnspan=2)

window.mainloop()
